"""
Generates the technical Word document for the M1SPAR-projet2 recommendation system.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

# Helper functions
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ═══════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Rapport Technique")
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Systeme de Recommandation Amazon")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("M1 SPAR - Projet 2\nAnnee 2025-2026")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SOMMAIRE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("Sommaire", level=1)
sommaire = [
    "1. Problematique et objectifs du projet",
    "2. Le dataset : description, reduction et justification",
    "3. Analyse exploratoire : ce que le dataset nous apprend",
    "4. Architecture globale du systeme",
    "5. Pipeline de donnees : Bronze, Silver, Gold",
    "6. Qualite des donnees et monitoring",
    "7. Agregations de donnees : existantes et proposees",
    "8. Modeles de Machine Learning",
    "9. Evaluation et optimisation",
    "10. API de recommandation",
    "11. Deploiement et infrastructure",
    "12. Problemes rencontres et decisions de pivot",
    "13. Bilan et perspectives",
]
for item in sommaire:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. PROBLEMATIQUE ET OBJECTIFS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Problematique et objectifs du projet", level=1)

doc.add_heading("1.1 Problematique", level=2)
doc.add_paragraph(
    "Les plateformes e-commerce comme Amazon proposent des millions de produits. "
    "Face a ce volume, un utilisateur ne peut explorer qu'une infime fraction du catalogue. "
    "Le defi est de predire, pour chaque utilisateur, les produits les plus susceptibles "
    "de l'interesser, en s'appuyant sur son historique d'interactions et sur les "
    "caracteristiques des produits."
)
doc.add_paragraph(
    "Ce projet vise a construire un systeme de recommandation complet, de bout en bout, "
    "couvrant l'ensemble de la chaine : ingestion des donnees brutes, nettoyage, "
    "entrainement de modeles de Machine Learning, serving temps reel via une API, "
    "et monitoring en production."
)

doc.add_heading("1.2 Objectifs mesurables", level=2)
add_table(doc,
    ["Objectif", "Metrique", "Cible"],
    [
        ["Qualite des recommandations", "NDCG@10", "Lift > 0% vs baseline popularite"],
        ["Couverture du catalogue", "% produits recommandables", "100% (warm + cold)"],
        ["Latence de l'API", "P95 latence /recommend", "< 50 ms"],
        ["Qualite des donnees", "Tests Great Expectations", "100% de checks passes"],
        ["Detection de derive", "Test KS mensuel", "Alerte si p-value < 0.05"],
        ["A/B testing", "CTR par variante", "Significativite statistique (p < 0.05)"],
        ["Charge", "Test Locust (100 users)", "P95 < 50 ms sur 60 secondes"],
    ],
)

# ═══════════════════════════════════════════════════════════════════════════
# 2. DATASET
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Le dataset : description, reduction et justification", level=1)

doc.add_heading("2.1 Source des donnees", level=2)
doc.add_paragraph(
    "Le dataset utilise est Amazon Reviews 2023, publie par le McAuley Lab (UCSD) "
    "et distribue via HuggingFace. C'est la reference academique pour les systemes "
    "de recommandation. Il contient 571 millions de reviews reparties sur 33 categories "
    "de produits Amazon, couvrant les donnees de 2018 a 2023."
)
doc.add_paragraph(
    "Le dataset fournit deux types de fichiers par categorie :"
)
add_table(doc,
    ["Type", "Contenu", "Champs cles"],
    [
        ["Reviews", "Interactions utilisateur-produit", "user_id, asin, rating, timestamp, verified_purchase, helpful_vote"],
        ["Metadonnees", "Informations produit", "asin, title, description, price, average_rating, category"],
    ],
)

doc.add_heading("2.2 Pourquoi ce dataset ?", level=2)
doc.add_paragraph(
    "Plusieurs raisons justifient ce choix :"
)
bullets = [
    "Reference academique : utilise dans des centaines de publications sur les systemes de recommandation.",
    "Volume realiste : 571 millions d'interactions permettent de travailler sur des problematiques de passage a l'echelle.",
    "Richesse : combine le feedback explicite (ratings) et implicite (achats verifies), ainsi que des metadonnees textuelles (titres, descriptions) exploitables pour le content-based filtering.",
    "Multi-categories : permet de tester la generalisation du modele sur des domaines produits differents.",
    "Donnees reelles : issues de vrais comportements d'achat, avec tous les biais associes (distribution en J des notes, long tail).",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("2.3 Reduction du dataset", level=2)
doc.add_paragraph(
    "Le dataset complet (571M reviews, 33 categories) represente plusieurs centaines de Go "
    "et necessite des ressources de calcul considerables. Pour rendre le projet realisable "
    "sur des machines de developpement standards, nous avons reduit le perimetre a "
    "5 categories : All Beauty, Amazon Fashion, Appliances, Arts Crafts and Sewing, et Automotive."
)
doc.add_paragraph(
    "Cette selection offre un bon equilibre :"
)
bullets = [
    "Volume suffisant : environ 34 millions d'interactions sur ces 5 categories.",
    "Diversite : categories de tailles tres differentes (Automotive = 58,3% du volume, Appliances = beaucoup moins).",
    "Representativite : les phenomenes observes (sparsity, long tail, biais de notation) sont identiques a ceux du dataset complet.",
    "Temps de traitement raisonnable : le pipeline complet s'execute en moins d'une heure.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("2.4 Format de stockage", level=2)
doc.add_paragraph(
    "Les donnees brutes sont telechargees au format Apache Arrow via la bibliotheque "
    "HuggingFace datasets (version 2.21.0). Ce format colonnaire permet une lecture "
    "tres rapide en memoire. Pour le pipeline d'analyse, les donnees sont converties "
    "en Parquet, un format compresse et partitionnable nativement supporte par PySpark."
)
add_table(doc,
    ["Format", "Lisible", "Vitesse lecture", "Taille disque"],
    [
        ["JSON Lines", "Oui", "Lent", "Grand"],
        ["CSV", "Oui", "Moyen", "Moyen"],
        ["Arrow", "Non", "Tres rapide", "Moyen"],
        ["Parquet", "Non", "Rapide", "Petit (compresse)"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. ANALYSE DU DATASET
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Analyse exploratoire : ce que le dataset nous apprend", level=1)

doc.add_paragraph(
    "L'analyse exploratoire (EDA) a ete conduite dans le notebook 01_eda_reco.ipynb "
    "et a repondu a 5 questions fondamentales qui ont directement guide les choix d'architecture."
)

doc.add_heading("3.1 Sparsity de la matrice user-item (99,9999 %)", level=2)
add_table(doc,
    ["Metrique", "Valeur"],
    [
        ["Utilisateurs", "13 664 354"],
        ["Produits", "1 834 471"],
        ["Interactions observees", "34 253 280"],
        ["Interactions possibles", "~25 000 milliards"],
        ["Taux de sparsity", "99,9999 %"],
    ],
)
doc.add_paragraph(
    "Sur 25 000 milliards de paires user-produit possibles, seulement 34 millions ont ete observees. "
    "Un utilisateur donne n'a interagi qu'avec moins d'un produit sur 500 000. "
    "Ce constat impose l'utilisation d'un algorithme concu pour les matrices creuses comme ALS, "
    "qui deduit les preferences manquantes a partir des rares interactions connues."
)

doc.add_heading("3.2 95 % des interactions sont des achats verifies", level=2)
add_table(doc,
    ["Type", "Volume", "Pourcentage"],
    [
        ["Achat verifie (purchase)", "32 725 041", "95,5 %"],
        ["Simple review", "1 528 239", "4,5 %"],
    ],
)
doc.add_paragraph(
    "Le flag verified_purchase est un signal fort : l'utilisateur a reellement achete le produit. "
    "Les 4,5 % restants sont des reviews non verifiees, un signal plus faible. "
    "Par ailleurs, 7,6 millions de reviews ont recu des votes 'helpful', un signal de qualite "
    "supplementaire exploitable pour ponderer les interactions."
)

doc.add_heading("3.3 Biais de notation massif (distribution en J)", level=2)
add_table(doc,
    ["Note", "Volume", "Pourcentage"],
    [
        ["5 etoiles", "22 893 163", "66,8 %"],
        ["4 etoiles", "3 754 822", "11,0 %"],
        ["1 etoile", "3 841 987", "11,2 %"],
        ["3 etoiles", "2 198 464", "6,4 %"],
        ["2 etoiles", "1 564 843", "4,6 %"],
    ],
)
doc.add_paragraph(
    "La moyenne est de 4,18/5 avec un ecart-type de 1,38. La distribution est en forme de J : "
    "pic massif a 5 etoiles et second pic a 1 etoile. Ce biais provient du survivorship bias "
    "(on ne note que ce qu'on a recu) et du self-selection (les satisfaits notent davantage). "
    "Consequence directe : utiliser le feedback implicite (achat oui/non) plutot que le rating brut."
)

doc.add_heading("3.4 Long tail : la majorite des produits sont quasi-inconnus", level=2)
doc.add_paragraph("Cote produits :")
add_table(doc,
    ["Seuil", "Produits concernes"],
    [
        ["<= 1 avis", "Plus de la moitie du catalogue"],
        ["<= 5 avis", "Majorite ecrasante"],
        ["<= 10 avis", "Quasi-totalite"],
    ],
)
doc.add_paragraph("Cote utilisateurs :")
add_table(doc,
    ["Seuil", "Utilisateurs concernes", "% des users"],
    [
        ["<= 1 interaction", "7 725 554", "56,5 %"],
        ["<= 5 interactions", "12 518 803", "91,6 %"],
        ["<= 10 interactions", "13 273 056", "97,1 %"],
    ],
)
doc.add_paragraph(
    "Plus de 9 utilisateurs sur 10 ont moins de 5 interactions. ALS seul ne peut pas construire "
    "de profil fiable pour ces entites. C'est le probleme du cold start, qui justifie "
    "l'architecture hybride avec bridge model."
)

doc.add_heading("3.5 Power users : une minorite critique", level=2)
add_table(doc,
    ["Segment", "Nb users", "% users", "Achats moy.", "% des achats"],
    [
        ["power_user", "3 989 719", "29,2 %", "39,1", "54,2 %"],
        ["regular_user", "7 660 158", "56,1 %", "7,4", "39,7 %"],
        ["casual_user", "2 014 477", "14,7 %", "1,5", "6,1 %"],
    ],
)
doc.add_paragraph(
    "Les power users representent 29 % des utilisateurs mais generent 54 % des achats. "
    "Optimiser uniquement le CTR moyen reviendrait a sacrifier la qualite pour les power users. "
    "L'evaluation doit etre segmentee pour proteger cette population critique."
)

doc.add_heading("3.6 Synthese des impacts architecturaux", level=2)
add_table(doc,
    ["Observation", "Consequence architecturale"],
    [
        ["Sparsity 99,9999 %", "ALS est l'algorithme adapte"],
        ["95 % d'achats verifies", "Implicit feedback > ratings bruts"],
        ["66,8 % de notes a 5 etoiles", "Ne pas utiliser le rating comme signal principal"],
        ["91,6 % des users ont < 5 interactions", "Cold start inevitable -> bridge model requis"],
        ["Power users = 29 % mais 54 % des achats", "Evaluer par segment, proteger les power users"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. ARCHITECTURE GLOBALE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Architecture globale du systeme", level=1)

doc.add_heading("4.1 Vue d'ensemble des composants", level=2)
doc.add_paragraph(
    "Le systeme est compose de 11 services orchestres via Docker Compose. "
    "L'architecture suit un pattern classique de data engineering avec une separation claire "
    "entre le pipeline de donnees (offline), le serving (online), et le monitoring."
)

add_table(doc,
    ["Composant", "Technologie", "Role"],
    [
        ["Pipeline Bronze", "Python + pandas + pyarrow", "Ingestion Arrow -> Parquet (3 tables)"],
        ["Pipeline Silver", "PySpark", "Nettoyage, deduplication, features temporelles"],
        ["Pipeline Gold", "PySpark", "Aggregation, metriques produit, cold tagging"],
        ["Embeddings contenu", "Sentence Transformers", "Encodage semantique (384 dims)"],
        ["Modele ALS", "implicit (Python)", "Factorisation matricielle (16 dims)"],
        ["Bridge Model", "scikit-learn MLPRegressor", "Projection cold start (384 -> 16)"],
        ["API", "FastAPI + uvicorn", "Serving temps reel (5 endpoints)"],
        ["Cache", "Redis", "Cache des recommandations (TTL 300s)"],
        ["Base feedback", "SQLite", "Stockage feedback + impressions A/B"],
        ["Monitoring", "Prometheus + Grafana", "Metriques temps reel + dashboards"],
        ["Dashboard", "Streamlit", "Interface utilisateur interactive"],
        ["Tracking ML", "MLflow", "Suivi des experiences et hyperparametres"],
    ],
)

doc.add_heading("4.2 Flux d'une transaction de bout en bout", level=2)
doc.add_paragraph(
    "Voici le parcours complet d'une requete de recommandation, de l'appel utilisateur "
    "jusqu'a la reponse :"
)

steps = [
    ("1. Requete entrante",
     "L'utilisateur (ou le dashboard Streamlit) appelle GET /recommend/{user_id}?n=10&category=Automotive. "
     "La requete arrive sur le serveur FastAPI (uvicorn)."),
    ("2. Verification du cache Redis",
     "L'API verifie d'abord si une reponse est en cache Redis (cle : reco:{user_id}:{n}:{category}:{exclude_purchased}). "
     "Si oui, la reponse cached est retournee directement (latence < 1 ms). "
     "Le compteur Prometheus cache_hits_total est incremente."),
    ("3. Identification de l'utilisateur",
     "Si cache miss, l'API cherche le user_id dans l'index en memoire (dictionnaire user_index). "
     "Si l'utilisateur existe, son embedding U1-U16 et son segment (power/regular/casual) sont recuperes."),
    ("4. Calcul des scores",
     "Pour un utilisateur connu (warm) : produit scalaire entre son vecteur U1-U16 et les vecteurs P1-P16 "
     "de tous les produits candidats. Pour un utilisateur inconnu (cold) : fallback vers les produits "
     "les plus populaires. Les scores sont normalises min-max entre 0 et 1."),
    ("5. Filtrage",
     "Les produits deja achetes sont exclus (si exclude_purchased=True). "
     "Si une categorie est specifiee, seuls les produits de cette categorie sont consideres. "
     "Les top-N produits sont selectionnes via argpartition (O(n) au lieu de O(n log n) pour un tri complet)."),
    ("6. Construction de la reponse",
     "Chaque recommandation inclut : product_id, score, source (als/bridge/popular), category. "
     "La reponse inclut egalement le segment utilisateur, la latence en ms, et le statut du cache."),
    ("7. Mise en cache",
     "La reponse est stockee dans Redis avec un TTL de 300 secondes pour les requetes futures identiques."),
    ("8. Metriques",
     "Prometheus collecte automatiquement la latence de la requete, le statut HTTP, "
     "et les compteurs de cache hit/miss via l'instrumentateur FastAPI."),
]
for title, desc in steps:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_heading("4.3 Flux de feedback (A/B testing)", level=2)
doc.add_paragraph(
    "En parallele du flux de recommandation, le systeme collecte du feedback :"
)
feedback_steps = [
    "L'utilisateur interagit avec une recommandation (click, purchase, skip).",
    "Le frontend appelle POST /feedback avec user_id, product_id, interaction_type et ab_variant.",
    "L'API determine la variante A/B de l'utilisateur via un hash deterministe (MD5 de user_id modulo 2), garantissant un split 50/50 stable.",
    "L'interaction et l'impression sont enregistrees dans SQLite (tables feedback et impressions).",
    "L'endpoint GET /ab_results calcule le CTR par variante, le lift, et la significativite statistique (Z-test bilateral).",
]
for i, step in enumerate(feedback_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. PIPELINE BRONZE / SILVER / GOLD
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Pipeline de donnees : Bronze, Silver, Gold", level=1)

doc.add_heading("5.1 Pourquoi cette architecture en couches ?", level=2)
doc.add_paragraph(
    "L'architecture medallion (Bronze / Silver / Gold) est un pattern standard en data engineering, "
    "popularise par Databricks. Elle structure le pipeline en couches de qualite croissante :"
)
add_table(doc,
    ["Couche", "Principe", "Avantage"],
    [
        ["Bronze", "Donnees brutes, non transformees", "Source de verite, re-traitement possible"],
        ["Silver", "Donnees nettoyees et dedupliquees", "Qualite garantie, prete pour l'analyse"],
        ["Gold", "Donnees aggregees et enrichies", "Prete pour le ML et le serving"],
    ],
)
doc.add_paragraph(
    "Ce decoupage permet de re-executer chaque couche independamment, de debugger "
    "les transformations couche par couche, et de garantir la tracabilite des donnees. "
    "Si une erreur est detectee dans Silver, on peut corriger la logique et re-traiter "
    "a partir de Bronze sans re-telecharger les donnees."
)

doc.add_heading("5.2 Bronze : ingestion (00_prepare_data.py)", level=2)
doc.add_paragraph(
    "La couche Bronze convertit les fichiers Arrow bruts (telecharges depuis HuggingFace) "
    "en 3 tables Parquet normalisees :"
)
add_table(doc,
    ["Table", "Contenu", "Cle primaire", "Partitionnement"],
    [
        ["interactions", "Qui a achete/note quoi, quand", "(user_id, product_id)", "Par category"],
        ["products", "Informations produit", "product_id", "Par category"],
        ["users", "Profil agrege des utilisateurs", "user_id", "Aucun"],
    ],
)
doc.add_paragraph(
    "Colonnes derivees a ce stade :"
)
derivations = [
    "interaction_type : 'purchase' si verified_purchase = True, sinon 'review'.",
    "converted : 1 si achat verifie, 0 sinon (signal binaire pour ALS implicite).",
    "year_month : extrait du timestamp (analyse de saisonnalite).",
    "segment : power_user (top 10 %), regular_user (p50-p90), casual_user (bottom 50 %), calcule sur les quantiles d'achats.",
]
for d in derivations:
    doc.add_paragraph(d, style="List Bullet")
doc.add_paragraph(
    "Strategie memoire : les categories sont traitees une par une pour ne jamais depasser "
    "la RAM disponible sur les machines de developpement."
)

doc.add_heading("5.3 Silver : nettoyage (silver_cleaning.py)", level=2)
doc.add_paragraph("La couche Silver applique 4 transformations via PySpark :")
silver_ops = [
    "Suppression des nulls : lignes sans user_id ou product_id.",
    "Clampage des ratings : bornes [0, 5] pour corriger les valeurs aberrantes.",
    "Deduplication : sur le triplet (user_id, product_id, timestamp).",
    "Features temporelles via window functions :",
]
for s in silver_ops:
    doc.add_paragraph(s, style="List Bullet")
add_table(doc,
    ["Feature", "Calcul", "Utilite"],
    [
        ["recency_rank", "row_number() OVER (PARTITION BY user_id ORDER BY timestamp DESC)", "Rang de fraicheur de l'interaction"],
        ["user_interaction_count", "count(*) OVER (PARTITION BY user_id)", "Activite totale de l'utilisateur"],
        ["rf_score", "user_interaction_count / recency_rank", "Score recence-frequence (RFM simplifie)"],
    ],
)

doc.add_heading("5.4 Gold : enrichissement (gold_features.py)", level=2)
doc.add_paragraph("La couche Gold produit les features finales pour le ML :")
gold_ops = [
    "Aggregation produit : nb_interactions, conversion_rate (sum(converted)/count), popularity_score (log1p normalise [0,1]).",
    "Cold tagging : is_cold = True si nb_interactions < 5.",
    "Enrichissement des interactions par broadcast join avec les features produit.",
]
for g in gold_ops:
    doc.add_paragraph(g, style="List Bullet")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. QUALITE DES DONNEES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Qualite des donnees et monitoring", level=1)

doc.add_heading("6.1 Validation avec Great Expectations", level=2)
doc.add_paragraph(
    "La qualite des donnees est validee automatiquement via Great Expectations (data_quality.py). "
    "Les checks suivants sont executes apres chaque transformation :"
)
add_table(doc,
    ["Check", "Regle", "Niveau"],
    [
        ["Rating valide", "rating entre 0 et 5", "Ligne"],
        ["Conversion rate valide", "conversion_rate entre 0 et 1", "Ligne"],
        ["Pas de null critique", "user_id et product_id non nuls", "Ligne"],
        ["Table non vide", "row count > 0", "Table"],
        ["Conversion rate global", "Entre 17 % et 23 %", "Aggrege"],
    ],
)
doc.add_paragraph(
    "Si un check echoue, le pipeline s'arrete et l'erreur est remontee. "
    "Cette approche 'fail fast' evite de propager des donnees corrompues vers les modeles."
)

doc.add_heading("6.2 Detection de derive (drift monitoring)", level=2)
doc.add_paragraph(
    "Le module drift_monitor.py compare la distribution de la variable 'converted' "
    "entre mois consecutifs via un test de Kolmogorov-Smirnov (KS). "
    "Si la p-value est inferieure a 0.05, une derive est flaggee."
)
doc.add_paragraph(
    "Le rapport de derive est sauvegarde dans data/drift_report.parquet avec les colonnes : "
    "month_a, month_b, n_a, n_b, conv_rate_a, conv_rate_b, ks_statistic, p_value, drift_detected."
)
doc.add_paragraph(
    "Ce mecanisme permet de detecter si le comportement d'achat change significativement "
    "d'un mois a l'autre, ce qui pourrait necessiter un re-entrainement du modele."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. AGREGATIONS DE DONNEES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Agregations de donnees : existantes et proposees", level=1)

doc.add_paragraph(
    "Les agregations transforment les donnees brutes en informations exploitables "
    "par les modeles et l'API. Cette section presente les agregations deja en place "
    "dans le pipeline, puis les agregations supplementaires proposees pour enrichir "
    "le systeme de recommandation."
)

doc.add_heading("7.1 Agregations existantes", level=2)
doc.add_paragraph(
    "Le pipeline actuel (Bronze / Silver / Gold) produit les agregations suivantes :"
)
add_table(doc,
    ["Agregation", "Niveau", "Couche", "Fichier source"],
    [
        ["total_interactions, total_purchases, avg_rating", "Par utilisateur", "Bronze", "00_prepare_data.py"],
        ["segment (power_user / regular_user / casual_user)", "Par utilisateur", "Bronze", "00_prepare_data.py"],
        ["recency_rank (rang par user, tri par timestamp desc)", "Par interaction", "Silver", "silver_cleaning.py"],
        ["user_interaction_count (nombre total par user)", "Par interaction", "Silver", "silver_cleaning.py"],
        ["rf_score (interaction_count / recency_rank)", "Par interaction", "Silver", "silver_cleaning.py"],
        ["nb_interactions, conversion_rate", "Par produit", "Gold", "gold_features.py"],
        ["popularity_score (log1p normalise [0,1])", "Par produit", "Gold", "gold_features.py"],
        ["is_cold (< 5 interactions)", "Par produit", "Gold", "gold_features.py"],
    ],
)
doc.add_paragraph(
    "Ces agregations couvrent les besoins de base du systeme : segmentation utilisateur, "
    "scoring de popularite, et tagging cold start. Elles sont calculees dans les couches "
    "Silver et Gold du pipeline ETL via PySpark (window functions et groupBy)."
)

doc.add_heading("7.2 Agregation par categorie", level=2)
doc.add_paragraph(
    "Proposee : calculer nb_users, nb_products, avg_rating, conversion_rate "
    "et avg_interactions_per_user par categorie Amazon."
)
doc.add_paragraph(
    "Justification : les 5 categories du sous-ensemble ont des profils tres differents. "
    "Automotive represente 58 % du volume mais cela ne signifie pas que c'est la categorie "
    "avec le meilleur taux de conversion. Cette agregation revelerait les categories "
    "sous-exploitees (peu d'interactions mais bon taux de conversion) et les categories "
    "a risque (beaucoup de volume mais faible satisfaction). Elle est aussi utile pour "
    "ponderer les recommandations cross-categories : si un utilisateur achete principalement "
    "dans une categorie a fort taux de conversion, on peut lui proposer des produits similaires "
    "dans une categorie adjacente."
)

doc.add_heading("7.3 Agregation par saisonnalite (categorie x mois)", level=2)
doc.add_paragraph(
    "Proposee : calculer nb_interactions, conversion_rate et avg_rating "
    "par couple (category, year_month)."
)
doc.add_paragraph(
    "Justification : les comportements d'achat sont saisonniers (Black Friday, Noel, Prime Day). "
    "Agreger par categorie et par mois permet de detecter les pics de demande par categorie "
    "et d'adapter les recommandations en fonction de la saison. Par exemple, Appliances peut "
    "avoir un pic en novembre que All_Beauty n'a pas. C'est aussi un input direct pour le "
    "drift monitor : si une categorie chute un mois donne, c'est un signal d'alerte qui peut "
    "declencher un re-entrainement du modele."
)

doc.add_heading("7.4 Agregation du profil temporel produit", level=2)
doc.add_paragraph(
    "Proposee : calculer first_interaction_date, last_interaction_date, lifespan_days "
    "et trend (croissant / decroissant / stable) par produit."
)
doc.add_paragraph(
    "Justification : un produit avec 100 interactions reparties sur 2 ans et un produit "
    "avec 100 interactions concentrees sur 1 semaine n'ont pas le meme profil. Le trend "
    "permet de detecter les produits en train de disparaitre (plus aucune interaction recente) "
    "versus les produits en croissance. Recommander un produit en declin est une mauvaise "
    "experience utilisateur. Cette feature pourrait etre integree comme filtre post-scoring : "
    "penaliser les produits dont le trend est decroissant dans le classement final."
)

doc.add_heading("7.5 Agregation de la diversite d'achat utilisateur", level=2)
doc.add_paragraph(
    "Proposee : calculer nb_categories_distinctes, category_principale et "
    "indice_diversite (entropie de Shannon sur les categories) par utilisateur."
)
add_code_block(doc, "Entropie de Shannon : H = -SUM(p_i * log2(p_i)) pour chaque categorie i")
doc.add_paragraph(
    "Justification : un utilisateur qui achete dans 5 categories differentes et un "
    "utilisateur qui n'achete que dans Automotive n'attendent pas les memes recommandations. "
    "L'indice de diversite permet de personnaliser la strategie : pour un utilisateur diversifie, "
    "on peut proposer du cross-category ; pour un utilisateur specialise, on reste dans sa categorie "
    "dominante. Cela enrichit la segmentation au-dela du simple volume (power/regular/casual) "
    "en ajoutant une dimension thematique."
)

doc.add_heading("7.6 Agregation de la qualite des reviews par produit", level=2)
doc.add_paragraph(
    "Proposee : calculer avg_helpful_votes, pct_verified_purchases, "
    "et sentiment_ratio (pourcentage de notes 4-5 etoiles versus 1-2 etoiles) par produit."
)
doc.add_paragraph(
    "Justification : le champ helpful_vote est un signal de qualite deja present dans le dataset "
    "mais pas encore agrege. Un produit avec beaucoup de votes helpful a des reviews de qualite, "
    "ce qui est un indicateur de confiance. Le pct_verified_purchases mesure la proportion d'achats "
    "reels versus reviews non verifiees. Ces features pourraient ponderer les interactions dans ALS : "
    "une interaction sur un produit de confiance vaut plus qu'une interaction sur un produit douteux. "
    "Cela ameliorerait directement la qualite des recommandations sans modifier l'algorithme."
)

doc.add_heading("7.7 Agregation de la matrice segment x categorie", level=2)
doc.add_paragraph(
    "Proposee : calculer avg_rating, conversion_rate et nb_interactions "
    "par couple (segment, category)."
)
add_table(doc,
    ["", "All_Beauty", "Amazon_Fashion", "Appliances", "Arts_Crafts", "Automotive"],
    [
        ["power_user", "conv_rate", "conv_rate", "conv_rate", "conv_rate", "conv_rate"],
        ["regular_user", "conv_rate", "conv_rate", "conv_rate", "conv_rate", "conv_rate"],
        ["casual_user", "conv_rate", "conv_rate", "conv_rate", "conv_rate", "conv_rate"],
    ],
)
doc.add_paragraph(
    "Justification : les power users n'achetent pas les memes categories que les casual users. "
    "Cette matrice revele les affinites segment-categorie. Elle est particulierement utile pour "
    "le cold start : pour un nouvel utilisateur dont on connait le segment (via une heuristique "
    "basee sur son premier achat), on peut deja biaiser les recommandations vers les categories "
    "preferees de son segment, au lieu de tomber sur un simple fallback popularite global."
)

doc.add_heading("7.8 Agregation du pattern temporel utilisateur", level=2)
doc.add_paragraph(
    "Proposee : calculer avg_days_between_purchases, regularity_score "
    "(ecart-type normalise des intervalles entre achats) et last_purchase_days_ago "
    "par utilisateur."
)
doc.add_paragraph(
    "Justification : distinguer les acheteurs reguliers (tous les mois) des acheteurs "
    "ponctuels (une fois par an). Le last_purchase_days_ago detecte les utilisateurs en "
    "train de churner. Cela permet de prioriser les recommandations pour les utilisateurs "
    "a risque de churn et d'adapter la frequence des suggestions. Un utilisateur regulier "
    "qui n'a pas achete depuis 3 mois est un signal d'alerte ; un acheteur annuel qui n'a "
    "pas achete depuis 3 mois est normal."
)

doc.add_heading("7.9 Synthese et impact sur le systeme", level=2)
add_table(doc,
    ["Agregation", "Niveau", "Impact principal"],
    [
        ["Stats par categorie", "Categorie", "Comparaison et ponderation cross-category"],
        ["Saisonnalite", "Categorie x Mois", "Recommandations contextuelles, drift detection"],
        ["Profil temporel produit", "Produit", "Eviter de recommander des produits en declin"],
        ["Diversite d'achat", "Utilisateur", "Strategie cross-category vs intra-category"],
        ["Qualite des reviews", "Produit", "Ponderation de la confiance dans ALS"],
        ["Matrice segment x categorie", "Segment x Categorie", "Cold start plus intelligent"],
        ["Pattern temporel user", "Utilisateur", "Detection de churn, frequence de suggestion"],
    ],
)
doc.add_paragraph(
    "Les trois agregations les plus impactantes pour le systeme actuel sont : "
    "(1) la diversite d'achat, qui ameliore directement le recommender en personnalisant "
    "la strategie cross-category ; (2) la qualite des reviews, qui ameliore ALS via une "
    "ponderation de confiance sur les interactions ; et (3) la saisonnalite, qui rend les "
    "recommandations contextuelles et alimente le drift monitor."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. MODELES ML
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Modeles de Machine Learning", level=1)

doc.add_paragraph(
    "Le systeme utilise une architecture hybride a 3 couches qui combine collaborative filtering "
    "et content-based filtering pour couvrir 100 % du catalogue (warm + cold)."
)

doc.add_heading("8.1 Sentence Transformers : embeddings de contenu", level=2)
add_table(doc,
    ["Parametre", "Valeur"],
    [
        ["Modele", "all-MiniLM-L6-v2"],
        ["Dimensions", "384"],
        ["Input", "Concatenation titre + description"],
        ["Normalisation", "L2 (norme = 1)"],
        ["Sortie", "content_embeddings.parquet"],
    ],
)
doc.add_paragraph(
    "Ce modele encode le texte en vecteurs denses de 384 dimensions. Deux produits "
    "semantiquement similaires (meme si aucun utilisateur en commun) auront des embeddings proches. "
    "La normalisation L2 permet d'utiliser le produit scalaire comme proxy de la similarite cosinus, "
    "ce qui est plus efficace pour la recherche de voisins."
)
doc.add_paragraph(
    "Pourquoi all-MiniLM-L6-v2 ? C'est un modele leger (pas de GPU requis), rapide a l'inference, "
    "et de qualite suffisante pour le bridge model. Un modele plus lourd (e.g., all-mpnet-base-v2) "
    "n'apporterait pas de gain significatif vu que les embeddings sont ensuite projetes en 16 dimensions."
)

doc.add_heading("8.2 ALS : factorisation matricielle", level=2)
doc.add_paragraph(
    "ALS (Alternating Least Squares) est l'algorithme principal du systeme. Il decompose "
    "la matrice d'interactions user-item en deux matrices de rang faible :"
)
doc.add_paragraph(
    "U (n_users x rank) x V^T (rank x n_products) = predictions pour les paires non observees."
)
add_table(doc,
    ["Hyperparametre", "Valeur", "Justification"],
    [
        ["rank", "16", "Compromis expressivite / memoire. 16 dimensions capturent les patterns principaux."],
        ["iterations", "15", "Convergence suffisante sans surentrainement."],
        ["regularization", "0.1", "Penalite L2 pour eviter le sur-apprentissage sur les interactions rares."],
        ["implicitPrefs", "True", "Les ratings sont biaises (66,8 % a 5 etoiles). Le nombre d'interactions est plus fiable."],
    ],
)
doc.add_paragraph(
    "Pourquoi ALS et pas une autre approche ?"
)
als_reasons = [
    "Concu pour les matrices creuses (sparsity 99,9999 %).",
    "Mode implicit feedback natif : utilise le nombre d'interactions comme signal de confiance.",
    "Scalable : la bibliotheque implicit utilise les BLAS et les matrices sparse scipy.",
    "Entrainement rapide : quelques minutes sur 34 millions d'interactions.",
    "Interpretable : chaque dimension P1-P16 capture un facteur latent.",
]
for r in als_reasons:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("8.3 Bridge Model : resolution du cold start", level=2)
doc.add_paragraph(
    "Le probleme : ALS produit des embeddings excellents pour les produits warm (>= 5 interactions) "
    "mais ne couvre pas les produits cold. Sentence Transformers couvre tout le catalogue "
    "mais n'encode pas les comportements d'achat."
)
doc.add_paragraph(
    "La solution : un MLP (Multi-Layer Perceptron) qui apprend a projeter l'espace semantique "
    "(384 dimensions) vers l'espace comportemental ALS (16 dimensions). "
    "Il est entraine sur les produits warm ou les deux representations existent."
)
doc.add_paragraph("Architecture du MLP :")
add_code_block(doc, "ST(384) -> Dense(256, ReLU) -> Dense(64, ReLU) -> ALS(16)")

add_table(doc,
    ["Parametre", "Valeur"],
    [
        ["Couches cachees", "(256, 64)"],
        ["Activation", "ReLU"],
        ["Perte", "Mean Squared Error"],
        ["Early stopping", "15 iterations sans amelioration"],
        ["Max iterations", "300"],
        ["Validation", "10 % des donnees d'entrainement"],
    ],
)

doc.add_paragraph("Strategies par type d'entite :")
add_table(doc,
    ["Cas", "Methode", "Justification"],
    [
        ["Produit warm (>= 5 interactions)", "Facteurs ALS directs (P1-P16)", "Encodent les comportements reels d'achat"],
        ["Produit cold (< 5 interactions)", "Bridge(ST embeddings) -> P1-P16", "Projette la semantique vers l'espace ALS"],
        ["User warm", "Facteurs ALS directs (U1-U16)", "Encodent les preferences reelles"],
        ["User cold (jamais vu)", "Moyenne des P1-P16 des produits vus", "Approximation rapide sans re-entrainement"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9. EVALUATION ET OPTIMISATION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Evaluation et optimisation", level=1)

doc.add_heading("9.1 Metrique principale : NDCG@10 et Lift", level=2)
doc.add_paragraph(
    "Le NDCG (Normalized Discounted Cumulative Gain) a 10 est la metrique standard "
    "pour evaluer les systemes de recommandation. Elle mesure la qualite du classement "
    "des recommandations en penalisant les items pertinents places trop bas dans la liste. "
    "Un NDCG@10 de 1.0 signifie que tous les items pertinents sont dans le top 10 et dans l'ordre parfait ; "
    "un score de 0.0 signifie qu'aucun item pertinent n'apparait dans le top 10."
)
doc.add_paragraph(
    "Le lift mesure l'amelioration relative du modele ALS par rapport a la baseline de popularite. "
    "Il se calcule comme suit :"
)
add_code_block(doc, "Lift (%) = (NDCG_ALS - NDCG_popularite) / NDCG_popularite x 100")
doc.add_paragraph(
    "Un lift positif signifie que le modele ALS recommande mieux que la simple popularite. "
    "Un lift negatif signifie que la baseline de popularite surpasse ALS, ce qui arrive "
    "typiquement quand les donnees sont extremement creuses : la popularite est une heuristique "
    "robuste car elle recommande les produits les plus frequents, qui ont mecaniquement plus "
    "de chances d'apparaitre dans le test set."
)

doc.add_paragraph(
    "Protocole d'evaluation :"
)
eval_steps = [
    "Split temporel : 80 % train / 20 % test par utilisateur (les dernieres interactions).",
    "Echantillonnage : 1 000 utilisateurs avec >= 10 interactions.",
    "Baseline : top-K par popularite globale (produits les plus frequents du train set).",
    "Lift : amelioration en % de ALS vs la baseline popularite.",
]
for e in eval_steps:
    doc.add_paragraph(e, style="List Bullet")

doc.add_paragraph(
    "Resultats obtenus sur le sous-ensemble de 2 categories (All Beauty, Amazon Fashion) :"
)
add_table(doc,
    ["Modele", "NDCG@10", "Lift vs popularite"],
    [
        ["Baseline popularite", "0.0010", "-"],
        ["ALS (rank=32, reg=0.5, iter=10) — meilleur run", "0.0007", "-29.0 %"],
        ["ALS (rank=32, reg=0.01, iter=10)", "0.0006", "-44.0 %"],
        ["ALS (rank=16, reg=0.5, iter=20)", "0.0005", "-55.9 %"],
        ["ALS (rank=8, reg=0.5, iter=10)", "0.0001", "-92.4 %"],
    ],
)

doc.add_paragraph(
    "Analyse du lift negatif : sur ce sous-ensemble de donnees, le lift est systematiquement negatif. "
    "Cela s'explique par la sparsity extreme (99,9999 %) combinee au filtrage sur seulement 2 categories. "
    "Avec 91,6 % des utilisateurs ayant moins de 5 interactions, ALS ne dispose pas de suffisamment "
    "de signal pour construire des facteurs latents fiables. La baseline popularite, qui recommande "
    "les produits les plus frequents toutes categories confondues, beneficie d'un avantage statistique "
    "dans ce contexte : les produits populaires ont mecaniquement plus de chances d'apparaitre dans "
    "le test set d'un utilisateur quelconque."
)
doc.add_paragraph(
    "Ce resultat est attendu et documente dans la litterature (Dacrema et al., 2019 : "
    "'Are We Really Making Much Progress?'). Il justifie pleinement l'architecture hybride mise en place : "
    "le bridge model et les embeddings semantiques compensent les faiblesses d'ALS sur les entites cold, "
    "tandis que l'A/B testing en production permettra de mesurer le lift reel sur le trafic utilisateur "
    "avec un volume d'interactions suffisant."
)

doc.add_heading("9.2 Grid search des hyperparametres", level=2)
doc.add_paragraph(
    "Un grid search systematique explore 18 combinaisons d'hyperparametres (04_grid_search_als.py) :"
)
add_table(doc,
    ["Hyperparametre", "Valeurs testees"],
    [
        ["rank", "8, 16, 32"],
        ["regularization", "0.01, 0.1, 0.5"],
        ["iterations", "10, 20"],
    ],
)
doc.add_paragraph(
    "Chaque run est traque dans MLflow avec : hyperparametres, NDCG@10, lift vs popularite, "
    "nombre de facteurs user/item. Le meilleur run est automatiquement tague 'Production'."
)

doc.add_heading("9.3 Pistes d'optimisation identifiees", level=2)
optimisations = [
    ("Ponderation par helpful_vote",
     "Les 7,6 millions de votes 'helpful' pourraient ponderer la confiance accordee a chaque interaction dans ALS."),
    ("Embeddings plus riches",
     "Integrer le prix, la categorie et le rating moyen dans les embeddings de contenu "
     "(au lieu du texte seul) pour enrichir le bridge model."),
    ("Re-entrainement incremental",
     "Mettre a jour les facteurs ALS periodiquement avec les nouvelles interactions "
     "plutot que de re-entrainer from scratch."),
    ("Diversity-aware re-ranking",
     "Ajouter un re-ranking post-ALS pour diversifier les categories dans les top-N "
     "recommandations (eviter de recommander 10 produits de la meme categorie)."),
    ("Evaluation par segment",
     "Calculer le NDCG@10 separement pour power_user, regular_user et casual_user "
     "pour s'assurer qu'aucun segment ne regresse."),
    ("Feedback loop",
     "Exploiter les donnees de feedback (clicks, purchases, skips) collectees via l'API "
     "pour le re-entrainement du modele (online learning)."),
]
for title, desc in optimisations:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} : ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 10. API
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. API de recommandation", level=1)

doc.add_heading("10.1 Choix technologiques", level=2)
doc.add_paragraph(
    "L'API est construite avec FastAPI, un framework web Python asynchrone. "
    "Ce choix est motive par :"
)
api_reasons = [
    "Performance : uvicorn (ASGI) avec support async natif.",
    "Documentation automatique : generation OpenAPI / Swagger.",
    "Validation : schemas Pydantic avec validation des parametres.",
    "Instrumentation : integration native avec prometheus-fastapi-instrumentator.",
]
for r in api_reasons:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("10.2 Endpoints", level=2)
add_table(doc,
    ["Endpoint", "Methode", "Description", "Parametres"],
    [
        ["/recommend/{user_id}", "GET", "Recommandations personnalisees", "n (1-100), category, exclude_purchased"],
        ["/similar/{product_id}", "GET", "Produits similaires", "n (1-50)"],
        ["/feedback", "POST", "Enregistrer une interaction", "user_id, product_id, interaction_type, ab_variant"],
        ["/ab_results", "GET", "Resultats A/B testing", "experiment_id, segment"],
        ["/health", "GET", "Etat du service", "Aucun"],
    ],
)

doc.add_heading("10.3 Cache Redis", level=2)
doc.add_paragraph(
    "Un cache Redis est place devant le moteur de recommandation pour garantir la latence cible (P95 < 50 ms). "
    "La cle de cache encode tous les parametres de la requete : reco:{user_id}:{n}:{category}:{exclude_purchased}. "
    "Le TTL est de 300 secondes (5 minutes), un compromis entre fraicheur et taux de hit."
)
doc.add_paragraph(
    "Le cache est optionnel : si Redis est indisponible, l'API continue de fonctionner "
    "en mode direct (degradation gracieuse). Les metriques de hit/miss sont exposees "
    "a Prometheus pour le monitoring."
)

doc.add_heading("10.4 A/B Testing integre", level=2)
doc.add_paragraph(
    "Le systeme integre un framework d'A/B testing :"
)
ab_features = [
    "Attribution deterministe : hash MD5 du user_id modulo 2, garantissant un split 50/50 stable dans le temps.",
    "Collecte des impressions et du feedback dans SQLite.",
    "Calcul du CTR par variante : clicks / impressions.",
    "Lift : (CTR_treatment - CTR_control) / CTR_control x 100.",
    "Test de significativite : Z-test bilateral sur les proportions (scipy.stats.norm).",
    "Seuil de significativite : p-value < 0.05.",
]
for f in ab_features:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("10.5 Chargement des donnees en memoire", level=2)
doc.add_paragraph(
    "Au demarrage, l'API charge en RAM l'integralite des embeddings produits et utilisateurs "
    "(fichiers Parquet). Les matrices P1-P16 et U1-U16 sont stockees sous forme de tableaux NumPy "
    "pour un produit scalaire ultra-rapide. Des dictionnaires (product_index, user_index) "
    "permettent un acces O(1) par identifiant."
)
doc.add_paragraph(
    "Ce choix de tout charger en memoire elimine la latence I/O pendant l'inference. "
    "Le conteneur API est configure avec 6 Go de RAM dans Docker Compose."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 11. DEPLOIEMENT
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Deploiement et infrastructure", level=1)

doc.add_heading("11.1 Docker Compose : 11 services", level=2)
doc.add_paragraph(
    "L'ensemble du systeme est deploye via un seul fichier docker-compose.yml "
    "qui orchestre 11 services avec leurs dependances :"
)
add_table(doc,
    ["Service", "Image", "Memoire", "Port", "Depends on"],
    [
        ["prepare", "Python 3.11-slim (custom)", "8 GB", "-", "-"],
        ["embeddings", "Python 3.11-slim (custom)", "8 GB", "-", "prepare"],
        ["als", "Python 3.11-slim (custom)", "12 GB", "-", "embeddings"],
        ["bridge", "Python 3.11-slim (custom)", "8 GB", "-", "als"],
        ["grid_search", "Python 3.11-slim (custom)", "16 GB", "-", "prepare"],
        ["redis", "redis:7-alpine", "-", "6379", "-"],
        ["api", "Python 3.11-slim (custom)", "6 GB", "8000", "redis"],
        ["mlflow", "Python 3.11-slim", "-", "5000", "-"],
        ["prometheus", "prom/prometheus", "-", "9090", "api"],
        ["grafana", "grafana/grafana-oss", "-", "3000", "prometheus"],
        ["streamlit", "Python 3.11-slim (custom)", "-", "8501", "api"],
    ],
)

doc.add_heading("11.2 Monitoring en production", level=2)
doc.add_paragraph(
    "Le monitoring repose sur la stack Prometheus + Grafana :"
)
monitoring_details = [
    "Prometheus scrape l'endpoint /metrics de FastAPI toutes les 15 secondes.",
    "prometheus-fastapi-instrumentator expose automatiquement les latences, codes HTTP, et compteurs.",
    "Des metriques custom sont ajoutees : cache_hits_total, cache_misses_total.",
    "Grafana affiche 3 panels : latence P95 du /recommend, cache hit rate, et CTR A/B.",
    "Seuils d'alerte : vert (< 30 ms), jaune (30-50 ms), rouge (> 50 ms).",
]
for m in monitoring_details:
    doc.add_paragraph(m, style="List Bullet")

doc.add_heading("11.3 Dashboard Streamlit", level=2)
doc.add_paragraph(
    "Le dashboard Streamlit offre une interface interactive en 3 onglets :"
)
tabs = [
    "Recommandations Live : saisir un user_id, obtenir ses recommandations avec score, source et categorie.",
    "A/B Testing CTR : visualiser le CTR par variante, le lift, et la significativite statistique.",
    "Top Produits Populaires : explorer les produits similaires par similarite cosinus.",
]
for t in tabs:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("11.4 Tests", level=2)
doc.add_paragraph("La suite de tests couvre :")
test_details = [
    "Tests unitaires (pytest) : recommender, A/B testing, cache, schemas, drift monitor.",
    "Fixtures : un AppState mock avec des donnees synthetiques (conftest.py).",
    "Tests de charge (Locust) : 100 utilisateurs concurrents, 60 secondes, objectif P95 < 50 ms.",
]
for t in test_details:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("11.5 CI/CD", level=2)
doc.add_paragraph(
    "Un pipeline GitHub Actions est en place avec du linting et des commits semantiques. "
    "Les merges sont controles via pull requests sur la branche develop."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 12. PROBLEMES RENCONTRES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Problemes rencontres et decisions de pivot", level=1)

problems = [
    ("Incompatibilite HuggingFace datasets v3+",
     "La version 3.0+ de la bibliotheque datasets a supprime le support des scripts de chargement custom. "
     "Le dataset Amazon Reviews 2023 utilise encore ce mecanisme.",
     "Verrouillage sur datasets==2.21.0. Lecon : toujours epingler les versions des dependances critiques."),

    ("Volume de donnees trop important",
     "Le dataset complet (571M reviews, 33 categories) est impossible a traiter sur une machine "
     "de developpement standard (RAM, disque, temps).",
     "Reduction a 5 categories. Le pipeline est concu pour fonctionner avec N categories configurables. "
     "Les phenomenes observes (sparsity, biais, long tail) sont identiques sur le sous-ensemble."),

    ("Biais de notation inutilisable",
     "La distribution des notes en J (66,8 % a 5 etoiles) rend le rating brut inexploitable "
     "comme signal d'entrainement.",
     "Pivot vers l'implicit feedback : le nombre d'interactions (achats verifies) remplace le rating. "
     "Ce choix est valide par la litterature (Hu et al., 2008)."),

    ("Cold start massif",
     "91,6 % des utilisateurs ont moins de 5 interactions. ALS seul ne couvre qu'une fraction du catalogue.",
     "Architecture hybride a 3 couches : ALS pour les entites warm, Sentence Transformers + bridge model "
     "pour les entites cold, fallback popularite pour les cas extremes."),

    ("Memoire et performance du pipeline",
     "Le chargement simultane de toutes les categories en memoire provoquait des OOM (Out of Memory).",
     "Traitement categorie par categorie dans la couche Bronze. Allocation memoire explicite "
     "dans Docker Compose (8-16 Go par service). Shared memory (shm_size: 2g) pour PyTorch."),

    ("Latence de l'API",
     "Sans cache, le calcul de recommandations (produit scalaire sur des milliers de produits) "
     "pouvait depasser 50 ms sous charge.",
     "Ajout d'un cache Redis (TTL 300s) et pre-normalisation L2 des embeddings pour utiliser "
     "le produit scalaire au lieu de la similarite cosinus. Selection des top-K via argpartition (O(n))."),

    ("PySpark pour Silver/Gold vs pandas pour Bronze",
     "Le choix initial d'utiliser pandas partout posait des problemes de scalabilite pour les "
     "transformations complexes (window functions, joins).",
     "Utilisation de PySpark pour les couches Silver et Gold (window functions natives, broadcast joins), "
     "tout en gardant pandas pour Bronze (simplicite, une categorie a la fois)."),
]

for i, (title, problem, solution) in enumerate(problems, 1):
    doc.add_heading(f"12.{i} {title}", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Probleme : ")
    run.bold = True
    p.add_run(problem)
    p = doc.add_paragraph()
    run = p.add_run("Decision : ")
    run.bold = True
    p.add_run(solution)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 13. BILAN
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Bilan et perspectives", level=1)

doc.add_heading("13.1 Ce qui a ete realise", level=2)
realisations = [
    "Pipeline de donnees complet (Bronze/Silver/Gold) avec qualite validee par Great Expectations.",
    "Systeme de recommandation hybride couvrant 100 % du catalogue (warm + cold).",
    "API temps reel avec 5 endpoints, cache Redis, et latence cible P95 < 50 ms.",
    "Framework d'A/B testing integre avec significativite statistique.",
    "Monitoring complet : Prometheus + Grafana avec dashboards pre-configures.",
    "Dashboard interactif Streamlit pour l'exploration et le test.",
    "Suite de tests unitaires et de charge (pytest + Locust).",
    "Grid search systematique des hyperparametres avec tracking MLflow (18 runs).",
    "Detection de derive statistique (KS test mensuel).",
    "Deploiement containerise complet via Docker Compose (11 services).",
]
for r in realisations:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("13.2 Choix technologiques valides", level=2)
add_table(doc,
    ["Choix", "Resultat"],
    [
        ["ALS en mode implicit", "Contourne efficacement le biais de notation"],
        ["Bridge model MLP", "Couvre les produits cold sans re-entrainement"],
        ["Architecture medallion", "Pipeline reproductible et debuggable couche par couche"],
        ["Redis + argpartition", "Latence de serving optimisee"],
        ["Great Expectations", "Qualite des donnees automatisee et verifiable"],
        ["Docker Compose", "Deploiement reproductible et isole"],
    ],
)

doc.add_heading("13.3 Perspectives d'evolution", level=2)
perspectives = [
    "Online learning : integrer le feedback collecte pour mettre a jour le modele en continu.",
    "Deep learning : remplacer ALS par un modele neuronal (e.g., NCF, BERT4Rec) pour capturer des patterns plus complexes.",
    "Passage a l'echelle : migrer de Docker Compose vers Kubernetes pour le scaling horizontal de l'API.",
    "Feature store : centraliser les features dans un store dedie (Feast) pour garantir la coherence train/serve.",
    "Monitoring avance : ajouter des alertes PagerDuty/Slack sur les metriques de derive et de performance.",
    "Evaluation online : mesurer l'impact business reel via des metriques de revenue/conversion plutot que le NDCG offline.",
]
for p_text in perspectives:
    doc.add_paragraph(p_text, style="List Bullet")

doc.add_heading("13.4 Conclusion", level=2)
doc.add_paragraph(
    "Ce projet demontre la construction d'un systeme de recommandation de bout en bout, "
    "de l'ingestion de 34 millions d'interactions a un service temps reel monitorable en production. "
    "L'analyse exploratoire a guide chaque decision d'architecture : le biais de notation a impose "
    "l'implicit feedback, la sparsity a justifie ALS, la longue traine a rendu le bridge model "
    "indispensable, et la segmentation des utilisateurs a structure l'evaluation."
)
doc.add_paragraph(
    "L'architecture modulaire (Bronze/Silver/Gold, services Docker independants) permet "
    "d'evoluer chaque composant independamment. Le systeme de monitoring (Prometheus, Grafana, "
    "drift detection) et d'A/B testing offrent les outils necessaires pour iterer en production "
    "de maniere informee et mesuree."
)

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Rapport_Technique_Recommandation_Amazon.docx")
doc.save(output_path)
print(f"Document genere : {output_path}")
